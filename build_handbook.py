"""Generate a proper DOCX handbook for the Forex Autopilot."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# Styles
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# â”€â”€ Title Page â”€â”€
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('AI Forex Autopilot')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('User Handbook & Command Reference')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x47, 0x5a, 0x6b)

doc.add_paragraph()
ver = doc.add_paragraph()
ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = ver.add_run('Version 1.0 â€” July 2026')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

doc.add_page_break()

# â”€â”€ Table of Contents â”€â”€
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Overview',
    '2. Quick Start',
    '3. Commands Reference',
    '4. Dashboard (Web UI)',
    '5. Strategy â€” Human S/R (H1+M15)',
    '6. Configuration (.env)',
    '7. Exness / MT5 Setup',
    '8. AI Setup (DeepSeek)',
    '9. Risk Management',
    '10. Trading Machine Setup',
    '11. Updating the Bot',
    '12. Safety & Warnings',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# â”€â”€ 1. Overview â”€â”€
doc.add_heading('1. Overview', level=1)
doc.add_paragraph(
    'AI Forex Autopilot is an autonomous forex/gold trading bot that combines '
    'technical analysis with optional AI decision support. It supports:'
)
bullets = [
    'Multiple brokers: Exness via MetaTrader 5 (recommended), Paper (local sim), or OANDA',
    'Multi-timeframe strategy: Daily bias, 4H structure, H1 levels, M15 entry',
    'AI market analysis via DeepSeek (or any OpenAI-compatible API)',
    'Hard risk management: position sizing, daily loss limits, drawdown protection',
    'Web dashboard for monitoring and control',
    'Telegram alerts for trades and events',
    'SQLite trade database with full history',
    'Backtesting on real MT5 historical data',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_paragraph(
    '\nâš ï¸ RISK WARNING: Trading forex can lose money fast. Default mode is PAPER. '
    'This is software, not financial advice.'
)

doc.add_page_break()

# â”€â”€ 2. Quick Start â”€â”€
doc.add_heading('2. Quick Start', level=1)
doc.add_paragraph('First time setup:')
doc.add_paragraph('1. Open PowerShell as Administrator', style='List Number')
doc.add_paragraph('2. Navigate to the project folder:', style='List Number')
code = doc.add_paragraph()
run = code.add_run('  cd C:\\Users\\Alex\\.openclaw\\workspace\\agents\\forge\\forex-autopilot')
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_paragraph('3. Activate Python environment:', style='List Number')
code = doc.add_paragraph()
run = code.add_run('  .venv\\Scripts\\activate')
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_paragraph('4. Install dependencies:', style='List Number')
code = doc.add_paragraph()
run = code.add_run('  pip install -r requirements.txt')
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_paragraph('5. Copy .env.example to .env and fill in your settings', style='List Number')
doc.add_paragraph('6. Start the bot:', style='List Number')
code = doc.add_paragraph()
run = code.add_run('  python main.py run')
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_page_break()

# â”€â”€ 3. Commands Reference â”€â”€
doc.add_heading('3. Commands Reference', level=1)

commands = [
    ('python main.py run', 'Start the trading loop (keeps running until stopped)'),
    ('python main.py run --once', 'Single tick (test mode)'),
    ('python main.py dashboard', 'Start the web dashboard on http://127.0.0.1:8787'),
    ('python main.py status', 'Show account equity, open trades, recent events'),
    ('python main.py mt5-test', 'Test Exness MT5 connection, symbols, and prices'),
    ('python main.py backtest --bars 500 --instrument EUR_USD', 'Simple backtest'),
    ('python main.py strategy-bt --instrument EUR_USD --months 3', 'Full strategy backtest (MTF)'),
    ('python main.py strategy-bt --instrument GBP_USD --months 3', 'Strategy backtest on GBP/USD'),
    ('python main.py strategy-bt --instrument XAU_USD --months 1', 'Strategy backtest on Gold (high risk)'),
]
for cmd, desc in commands:
    p = doc.add_paragraph()
    run = p.add_run(cmd)
    run.bold = True
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    p2 = doc.add_paragraph(desc)
    p2.paragraph_format.space_after = Pt(6)
    p2.paragraph_format.space_before = Pt(0)

doc.add_paragraph()
doc.add_paragraph('To stop the bot:')
stops = [
    'Dashboard: click STOP BOT button',
    'CLI: create file data\\bot.stop (or delete data\\bot.pid)',
    'Keyboard: Ctrl+C in the terminal',
]
for s in stops:
    doc.add_paragraph(s, style='List Bullet')

doc.add_page_break()

# â”€â”€ 4. Dashboard â”€â”€
doc.add_heading('4. Dashboard (Web UI)', level=1)
doc.add_paragraph(
    'The dashboard runs on http://127.0.0.1:8787 and auto-refreshes every 15 seconds. '
    'Keep the dashboard terminal window open â€” use a second terminal for the trading bot.'
)
doc.add_paragraph('Dashboard tabs:')
tabs = [
    ('Overview', 'Account equity/balance, open trades, recent closed trades, AI insights'),
    ('Strategy', 'Change strategy preset with description'),
    ('Settings', 'Adjust risk, instruments, AI provider, SL/TP pips â€” changes take effect on next tick'),
    ('Log', 'Recent event log with timestamps'),
    ('Backtest', 'Run strategy backtests from the UI (runs in background)'),
]
for name, desc in tabs:
    p = doc.add_paragraph()
    run = p.add_run(f'â€¢ {name}: ')
    run.bold = True
    p.add_run(desc)

doc.add_paragraph()
doc.add_paragraph('TIP: Open two PowerShell windows:')
docs = ['Window 1: python main.py dashboard', 'Window 2: python main.py run']
for d in docs:
    doc.add_paragraph(d, style='List Bullet')

doc.add_page_break()

# â”€â”€ 5. Strategy â”€â”€
doc.add_heading('5. Strategy â€” Human S/R (H1+M15)', level=1)
doc.add_paragraph(
    'Name: Human Support & Resistance â€” Multi-Timeframe Approach'
)
doc.add_paragraph(
    'This strategy uses 4 timeframes to find high-probability entries:'
)

# Strategy table
table = doc.add_table(rows=5, cols=2)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

cells = [
    ('Timeframe', 'Role'),
    ('Daily (D)', 'MACRO BIAS FILTER â€” never trade against the daily trend'),
    ('4H', 'MEDIUM STRUCTURE â€” confirms trend, major swing S/R'),
    ('H1', 'KEY LEVELS â€” entry S/R zones from swing highs/lows'),
    ('M15', 'EXECUTION â€” wait for candle reaction at H1 level, then enter'),
]
for i, (col1, col2) in enumerate(cells):
    row = table.rows[i]
    row.cells[0].text = col1
    row.cells[1].text = col2
    if i == 0:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Best for: ')
run.bold = True
p.add_run('EUR/USD âœ…  |  GBP/USD âœ…')
p2 = doc.add_paragraph()
run = p2.add_run('High risk: ')
run.bold = True
p2.add_run('XAU/USD (Gold) âš ï¸ â€” high volatility and drawdown')

doc.add_page_break()

# â”€â”€ 6. Configuration â”€â”€
doc.add_heading('6. Configuration (.env file)', level=1)
doc.add_paragraph('Key settings in .env (edit with Notepad or any text editor):')

settings_list = [
    ('TRADING_MODE', 'paper | practice | live'),
    ('BROKER', 'paper | mt5 | oanda'),
    ('INSTRUMENTS', 'EUR_USD,GBP_USD,XAU_USD'),
    ('STRATEGY_ID', 'human_sr_h1_m15 (default)'),
    ('RISK_PER_TRADE_PCT', '0.75 (0.75% risk per trade)'),
    ('DAILY_LOSS_LIMIT_PCT', '3.0 (stop after 3% daily loss)'),
    ('MAX_DRAWDOWN_PCT', '8.0 (halt at 8% drawdown)'),
    ('MAX_OPEN_TRADES', '2 (max concurrent)'),
    ('LOOP_SECONDS', '60 (check every 60s)'),
]
table = doc.add_table(rows=len(settings_list)+1, cols=2)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.rows[0].cells[0].text = 'Setting'
table.rows[0].cells[1].text = 'Values'
for row in table.rows[0].cells:
    for p in row.paragraphs:
        for r in p.runs:
            r.bold = True
for i, (k, v) in enumerate(settings_list, 1):
    table.rows[i].cells[0].text = k
    table.rows[i].cells[1].text = v

doc.add_page_break()

# â”€â”€ 7. Exness / MT5 Setup â”€â”€
doc.add_heading('7. Exness / MT5 Setup', level=1)
doc.add_paragraph('Follow these steps to connect the bot to your Exness demo account:')

doc.add_heading('Step 1: Install & Login', level=2)
steps = [
    'Install MetaTrader 5 from Exness website',
    'Log into your Exness demo account inside MT5',
    'Note from MT5 / Exness cabinet: Login number, Trading password, Server name',
]
for s in steps:
    doc.add_paragraph(s, style='List Number')

doc.add_heading('Step 2: Enable Algo Trading', level=2)
doc.add_paragraph('In MT5:')
items = [
    'Click the Algo Trading button in the toolbar (must be green/on)',
    'Tools â†’ Options â†’ Expert Advisors',
    'Check "Allow algorithmic trading"',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Step 3: Configure .env', level=2)
code = doc.add_paragraph()
run = code.add_run(
    'TRADING_MODE=practice\n'
    'BROKER=mt5\n'
    'MT5_LOGIN=123456789\n'
    'MT5_PASSWORD=your_trading_password\n'
    'MT5_SERVER=Exness-MT5Trial11\n'
    'MT5_SYMBOL_MAP=EUR_USD:EURUSDm,GBP_USD:GBPUSDm,XAU_USD:XAUUSDm'
)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_heading('Step 4: Test Connection', level=2)
code = doc.add_paragraph()
run = code.add_run('python main.py mt5-test')
run.font.name = 'Consolas'
run.font.size = Pt(10)
doc.add_paragraph('You should see balance, equity, and resolved symbol prices.')

doc.add_heading('Step 5: Run on Demo', level=2)
code = doc.add_paragraph()
run = code.add_run(
    'python main.py run --once    # test one tick\n'
    'python main.py run            # start trading loop'
)
run.font.name = 'Consolas'
run.font.size = Pt(9)
doc.add_paragraph('Only switch to live after demo results look acceptable.')

doc.add_page_break()

# â”€â”€ 8. AI Setup â”€â”€
doc.add_heading('8. AI Setup (DeepSeek)', level=1)
doc.add_paragraph(
    'The AI provides second opinions on trades â€” it reviews the chart and either agrees '
    'or cautions. Without an API key, the bot uses pure technical analysis.'
)
doc.add_paragraph('Recommended: DeepSeek')

code = doc.add_paragraph()
run = code.add_run(
    'AI_PROVIDER=openai\n'
    'OPENAI_API_KEY=sk-your-key-here\n'
    'OPENAI_BASE_URL=https://api.deepseek.com/v1\n'
    'OPENAI_MODEL=deepseek-chat'
)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_paragraph()
doc.add_paragraph('Alternative: OpenAI')
code = doc.add_paragraph()
run = code.add_run(
    'AI_PROVIDER=openai\n'
    'OPENAI_API_KEY=sk-your-key-here\n'
    'OPENAI_BASE_URL=https://api.openai.com/v1\n'
    'OPENAI_MODEL=gpt-4o-mini'
)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_page_break()

# â”€â”€ 9. Risk â”€â”€
doc.add_heading('9. Risk Management', level=1)
doc.add_paragraph('Default risk parameters (all editable in .env or dashboard):')
risks = [
    '0.75% equity risk per trade',
    '3% daily loss kill switch',
    '8% max drawdown halt',
    'Max 2 open trades at once',
    'Max 8 trades per day',
    'Minimum 1.5 R:R (reward-to-risk)',
    'Cooldown after 3 consecutive losses (60 min)',
    'London + New York trading sessions only',
]
for r in risks:
    doc.add_paragraph(r, style='List Bullet')

doc.add_paragraph(
    '\nThe bot tracks peak equity and halts if drawdown exceeds the limit. '
    'When halted, it will not open new trades until manually reset.'
)

doc.add_page_break()

# â”€â”€ 10. Trading Machine Setup â”€â”€
doc.add_heading('10. Trading Machine Setup (Second PC)', level=1)
doc.add_paragraph('To run the bot on a dedicated machine:')
steps = [
    'Install Python 3.12 from python.org',
    'Install MetaTrader 5 + log into Exness demo',
    'Enable Algo Trading in MT5 (green button)',
    'Install Git from git-scm.com',
    'Open CMD as Administrator',
    'Clone the repo: git clone https://github.com/ethanadeltd/forex-autopilot.git',
    'cd forex-autopilot',
    'python -m venv .venv && .venv\\Scripts\\activate',
    'pip install -r requirements.txt',
    'copy .env.example .env (fill in credentials)',
    'python main.py mt5-test (verify connection)',
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {step}', style='List Number')

doc.add_paragraph()
doc.add_paragraph('Two terminals needed:')
code = doc.add_paragraph()
run = code.add_run('# Terminal 1: Bot\npython main.py run\n\n# Terminal 2: Dashboard\npython main.py dashboard --host 0.0.0.0 --port 8787')
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_paragraph()
doc.add_paragraph('For remote access, use ngrok:')
code = doc.add_paragraph()
run = code.add_run('winget install ngrok\nngrok http 8787')
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_page_break()

# â”€â”€ 11. Updating the Bot â”€â”€
doc.add_heading('11. Updating the Bot', level=1)
doc.add_paragraph(
    'When new features or fixes are pushed to GitHub, update the trading machine with a single command:'
)

doc.add_heading('On the Trading Machine:', level=2)
update_steps = [
    'Open CMD or PowerShell',
    'cd C:\\forex-autopilot',
    '.venv\\Scripts\\activate',
    'git pull  (downloads the latest code)',
    'pip install -r requirements.txt  (if dependencies changed)',
    'python main.py run  (restart the bot)',
]
for s in update_steps:
    doc.add_paragraph(s, style='List Number')

doc.add_paragraph()
code = doc.add_paragraph()
run = code.add_run(
    '# Full update sequence:\n'
    'cd C:\\forex-autopilot\n'
    '.venv\\Scripts\\activate\n'
    'git pull\n'
    'pip install -r requirements.txt\n'
    'python main.py run'
)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_paragraph()
doc.add_paragraph('On the developer machine (Alex\'s PC), the workflow is:', style='List Bullet')
dev_steps = [
    'Make changes to the code',
    'git add . && git commit -m "description of changes"',
    'git push',
    'Trading machine: git pull (see above)',
    'Restart the bot',
]
for s in dev_steps:
    doc.add_paragraph(s, style='List Number')

doc.add_page_break()

# â”€â”€ 12. Safety â”€â”€
doc.add_heading('12. Safety & Warnings', level=1)
doc.add_paragraph('âš ï¸ Important safety information:')

warnings = [
    'Default mode is PAPER â€” no real money at risk',
    'Test on demo (practice) for weeks before considering live',
    'Never risk money you cannot afford to lose',
    'Keep MT5 running AND Algo Trading enabled',
    'Symbol names differ by account â€” use MT5_SYMBOL_MAP',
    'XAU/USD (Gold) is high risk with this strategy',
    'Monitor the bot regularly â€” don\'t leave unattended for long periods',
    'Always verify stop-loss and take-profit levels',
    'This is software, not financial advice',
]
for w in warnings:
    doc.add_paragraph(w, style='List Bullet')

doc.add_paragraph()

# Footer
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('â€” End of Handbook â€”')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

# Save
output_path = 'C:\\Users\\Alex\\Desktop\\Forex Autopilot Handbook.docx'
doc.save(output_path)
print("Handbook saved to: %s" % output_path)
